from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate

import streamlit as st
import os
from docx import Document
from docx.shared import Inches
import io

from PIL import Image
import requests

#load the model
def load_llm(max_tokens, prompt_template):
    # Load the locally downloaded model here
    llm = CTransformers(
        model = "llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens = max_tokens,
        temperature = 0.7
    )
    llm_chain = LLMChain(
        llm = llm,
        prompt = PromptTemplate.from_template(prompt_template)
    )
    return llm
    
    import requests

def fetch_photo(query):
    api_key = 'kWmCxiNE7hh1uiLSsyKGi110yApV0mEGMFON10LDXbb9zp3bt4X0Zlfp'

    url =  'https://api.pexels.com/v1/search'

    headers = {
        'Authorization': api_key
    }

    params = {
        'query': query,
        'per_page': 1
    }

    response = requests.get(url, headers=headers, params=params)

    #Check if the request was successful (status code is 200)

    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            src_original_url = photos[0]['src']['original']
            return src_original_url
        else:
            print("No photos found!")

    else:
        print(f"Error:  {response.status_code}, {response.text}")

        return None

def create_word_docx(user_input, paragraph, image_input):
    doc = Document()

    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)

    doc.add_heading('Image', level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_input.seek(0)
    doc.add_picture(image_stream, width=Inches(4))

    return doc

st.set_page_config(layout="wide")

def main():

    st.title('Article Generator app using Llama 2')
    user_input = st.text_input("Please enter your article topic or idea for article generation")
    image_input = st.text_input("Please enter your topic or idea for image generation")
    
    if len(user_input) and len(image_input) > 0:

        col1, col2, col3 = st.columns([1, 2, 1])

        with col1:
            st.subheader("Generated Content by Llama 2")
            st.write("Your submitted idea/topic for article generation: " + user_input)
            st.write("Your submitted idea/topic for image generation: " + image_input )

            prompt_template = """"You are a digital marketing and SEO expert and your task is to generate articles for a given topice. So write an article on {topic} under 800 words. Your article must be length of 800 words. Stick to the topic given by the user and maintain a professional but at the same time creative tone."""   

            llm_call = load_llm(max_tokens=800, prompt_template=prompt_template)
            print(llm_call)
            result = llm_call(user_input)

            if len(result) > 0:
                st.info("Your content has been generated successfully")
                st.write(result)

            else:
                st.error("Sorry, we couldn't generate content for this topic")

        with col2:
            st.subheader("Your article's Fetched Image")
            image_url = "temp_image.jpg"
            st.image(image_url)

        with col3:
            st.subheader("Download Your Article")
            #image_input = "temp_image.jpg"
            image_response = requests.get(image_url)
            img = Image.open(io.BytesIO(image_response.content))
            doc = create_word_docx(user_input, result['text'], img)

            #save the word doc to BytesIO Buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            #streamlit download button
            st.download_button(
                label="Download Article",
                data = doc_buffer,
                file_name = "doc_final.docx",
                mime= "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


if __name__ == "__main__":
    main()